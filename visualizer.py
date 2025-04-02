from flask import Flask, render_template, request, jsonify
import spacy
from spacy import displacy
import fitz
from docx import Document
from docx.shared import RGBColor
from io import BytesIO
from collections import Counter

app = Flask(__name__)
nlp = spacy.load("./output/model-best")

if "parser" not in nlp.pipe_names and "sentencizer" not in nlp.pipe_names:
    nlp.add_pipe("sentencizer")

COLORS = {
    "PERSON": "lime",
    "LIT_WORK": "blue",
    "ART_WORK": "green",
    "ART_MOVEMENT": "orange",
    "ORG": "purple",
    "PLACE": "teal",
    "EVENT": "cyan",
    "GENRE": "pink",
    "CHARACTER": "brown",
    "QUOTE": "gold",
    "AWARD": "lime",
    "PERIOD": "magenta",
    "TECHNIQUE": "indigo",
    "MOVIE_TV": "violet"
}

ENTITY_COLORS = {
    "PERSON": RGBColor(50, 205, 50),         # Lime
    "LIT_WORK": RGBColor(0, 0, 255),         # Blue
    "ART_WORK": RGBColor(0, 128, 0),         # Green
    "ART_MOVEMENT": RGBColor(255, 165, 0),     # Orange
    "ORG": RGBColor(128, 0, 128),            # Purple
    "PLACE": RGBColor(0, 128, 128),          # Teal
    "EVENT": RGBColor(0, 255, 255),          # Cyan
    "GENRE": RGBColor(255, 20, 147),         # Pink
    "CHARACTER": RGBColor(165, 42, 42),      # Brown
    "QUOTE": RGBColor(255, 215, 0),          # Gold
    "AWARD": RGBColor(50, 205, 50),          # Lime
    "PERIOD": RGBColor(255, 0, 255),         # Magenta
    "TECHNIQUE": RGBColor(75, 0, 130),       # Indigo
    "MOVIE_TV": RGBColor(148, 0, 211)        # Violet
}

# Get the entity labels from the model's NER pipe so its not dependant on the doc
MODEL_ENTITIES = sorted(nlp.get_pipe("ner").labels)

@app.route("/")
def index():
    entity_types = MODEL_ENTITIES
    doc = nlp("")
    options = {
        "ents": entity_types,
        "colors": {ent: COLORS.get(ent, "gray") for ent in entity_types}
    }
    html = displacy.render(doc, style="ent", options=options, page=False)
    return render_template("index.html", entity_types=entity_types, html=html)


@app.route("/filter", methods=["POST"])
def filter_entities():
    """
    Filter and highlight entities based on the selected types.
    """
    data = request.json
    selected_entities = data.get("selected_entities", [])
    user_text = data.get("text", "").strip()

    doc = nlp(user_text)
    options = {
        "ents": selected_entities,
        "colors": {ent: COLORS.get(ent, "gray") for ent in selected_entities}
    }
    html = displacy.render(doc, style="ent", options=options, page=False)
    return jsonify({"html": html})


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text("text") + "\n"
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    text = "\n".join(para.text for para in doc.paragraphs)
    return text.strip()


@app.route("/upload", methods=["POST"])
def upload_file():
    """
    Handle file uploads and extract text from PDF or DOCX files.
    """
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    filename = file.filename.lower()
    if not (filename.endswith(".pdf") or filename.endswith(".docx")):
        return jsonify({"error": "Invalid file format. Please upload a PDF or DOCX."}), 400

    file_bytes = file.read()
    if filename.endswith(".pdf"):
        extracted_text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        extracted_text = extract_text_from_docx(file_bytes)

    return jsonify({"text": extracted_text})


@app.route("/save", methods=["POST"])
def save_results():
    """
    Save the processed text with highlighted entities to a DOCX file.
    """
    data = request.json
    text = data.get("text", "").strip()
    selected_entities = data.get("selected_entities", [])

    if not text:
        return jsonify({"error": "No content to save"}), 400

    doc = Document()
    para = doc.add_paragraph()
    processed_doc = nlp(text)

    last_index = 0
    for ent in processed_doc.ents:
        if ent.label_ in selected_entities:
            if ent.start_char > last_index:
                para.add_run(text[last_index:ent.start_char])
            entity_run = para.add_run(ent.text)
            entity_run.bold = True
            entity_run.font.color.rgb = ENTITY_COLORS.get(
                ent.label_, RGBColor(128, 128, 128)
            )
            last_index = ent.end_char

    if last_index < len(text):
        para.add_run(text[last_index:])

    file_path = "Results.docx"
    doc.save(file_path)
    return jsonify({"success": True, "file_path": file_path})

@app.route("/stats", methods=["POST"])
def stats():
    data = request.json
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "No text provided"}), 400

    doc = nlp(text)

    # Entity Count & Distribution
    entity_counts = Counter(ent.label_ for ent in doc.ents)
    total_entities = len(doc.ents)

    # Token Count
    token_count = len(doc)

    # Entity Density (entities per token)
    entity_density = total_entities / token_count if token_count > 0 else 0

    # Sentence Length and Structure (list and average)
    sentence_lengths = [len(sent) for sent in doc.sents]
    avg_sentence_length = sum(sentence_lengths) / len(sentence_lengths) if sentence_lengths else 0

    # Named Entity Length Distribution (token count for each entity)
    ne_length_distribution = [len(ent) for ent in doc.ents]

    return jsonify({
        "entity_count": total_entities,
        "entity_distribution": dict(entity_counts),
        "token_count": token_count,
        "entity_density": entity_density,
        "sentence_lengths": sentence_lengths,
        "avg_sentence_length": avg_sentence_length,
        "named_entity_length_distribution": ne_length_distribution,
    })

if __name__ == "__main__":
    app.run(debug=True)
